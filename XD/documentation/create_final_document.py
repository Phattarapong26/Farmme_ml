#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
สร้างเอกสารวิชาการ FarmMe ฉบับสมบูรณ์
รวมทุกบทและภาคผนวก
"""

import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Import functions from part 2
sys.path.insert(0, str(Path(__file__).parent))
from doc_generator_part2 import add_chapter2, add_chapter3, add_chapter4, add_chapter5

def main():
    print("="*80)
    print("🚀 สร้างเอกสารวิชาการ FarmMe ฉบับสมบูรณ์")
    print("="*80)
    
    # โหลดเอกสารที่สร้างไว้แล้ว
    doc_path = Path(__file__).parent / "เอกสารวิชาการ_FarmMe_Complete.docx"
    
    if doc_path.exists():
        print(f"\n📂 โหลดเอกสารเดิม: {doc_path}")
        doc = Document(str(doc_path))
    else:
        print("\n❌ ไม่พบเอกสารเดิม กรุณารัน generate_full_document.py ก่อน")
        return
    
    # เพิ่มบทที่ 2-5
    print("\n📄 เพิ่มบทที่ 2: ทฤษฎีและงานวิจัยที่เกี่ยวข้อง...")
    add_chapter2(doc)
    
    print("📄 เพิ่มบทที่ 3: การดำเนินงาน...")
    add_chapter3(doc)
    
    print("📄 เพิ่มบทที่ 4: ผลลัพธ์...")
    add_chapter4(doc)
    
    print("📄 เพิ่มบทที่ 5: สรุป...")
    add_chapter5(doc)
    
    # เพิ่มบรรณานุกรม
    print("📄 เพิ่มบรรณานุกรม...")
    add_bibliography(doc)
    
    # เพิ่มภาคผนวก
    print("📄 เพิ่มภาคผนวก...")
    add_appendix(doc)
    
    # บันทึกไฟล์
    output_path = Path(__file__).parent / "เอกสารวิชาการ_FarmMe_ฉบับสมบูรณ์.docx"
    doc.save(str(output_path))
    
    print("\n" + "="*80)
    print("✅ สร้างเอกสารสำเร็จ!")
    print("="*80)
    print(f"📁 ไฟล์: {output_path}")
    print(f"📄 ขนาดไฟล์: {output_path.stat().st_size / 1024:.2f} KB")
    print("\n🎉 เอกสารวิชาการฉบับสมบูรณ์พร้อมใช้งาน!")

def add_bibliography(doc):
    """เพิ่มบรรณานุกรม"""
    doc.add_heading('บรรณานุกรม', level=1)
    
    references = [
        'Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. In Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining (pp. 785-794).',
        'Chapelle, O., & Li, L. (2011). An empirical evaluation of thompson sampling. In Advances in neural information processing systems (pp. 2249-2257).',
        'Deb, K., Pratap, A., Agarwal, S., & Meyarivan, T. A. M. T. (2002). A fast and elitist multiobjective genetic algorithm: NSGA-II. IEEE transactions on evolutionary computation, 6(2), 182-197.',
        'Liakos, K. G., Busato, P., Moshou, D., Pearson, S., & Bochtis, D. (2018). Machine learning in agriculture: A review. Sensors, 18(8), 2674.',
        'Kamilaris, A., & Prenafeta-Boldú, F. X. (2018). Deep learning in agriculture: A survey. Computers and electronics in agriculture, 147, 70-90.',
        'Wolfert, S., Ge, L., Verdouw, C., & Bogaardt, M. J. (2017). Big data in smart farming–a review. Agricultural Systems, 153, 69-80.',
    ]
    
    for ref in references:
        doc.add_paragraph(ref, style='List Number')
    
    doc.add_page_break()

def add_appendix(doc):
    """เพิ่มภาคผนวก"""
    doc.add_heading('ภาคผนวก', level=1)
    
    # ภาคผนวก ก
    doc.add_heading('ภาคผนวก ก: โค้ดตัวอย่าง', level=2)
    
    code_example = """
# ตัวอย่างการใช้งาน Model A
from Model_A_Fixed.model_algorithms_clean import ModelA_XGBoost

model = ModelA_XGBoost()
model.train(X_train, y_train)
predictions = model.predict(X_test)

# ตัวอย่างการใช้งาน Pipeline
from Pipeline_Integration.pipeline import FarmingPipeline

pipeline = FarmingPipeline(
    farmer_id='F001',
    farm_size_rai=25,
    budget_baht=150000
)

# Stage 1: Crop Selection
crops = pipeline.stage_1_crop_selection(model_a_results)

# Stage 2: Planting Window
window = pipeline.stage_2_planting_window(model_b_result)

# Stage 3: Price Forecast
forecast = pipeline.stage_3_price_forecast(model_c_result, dates)

# Stage 4: Harvest Decision
decision = pipeline.stage_4_harvest_decision(model_d_result, price, yield_kg)
    """
    
    p = doc.add_paragraph(code_example)
    p.style = 'Normal'
    
    # ภาคผนวก ข
    doc.add_heading('ภาคผนวก ข: ตารางข้อมูล', level=2)
    
    doc.add_paragraph('ตารางที่ 1: สรุปชุดข้อมูลที่ใช้ในโครงการ')
    
    # สร้างตาราง
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header
    header_cells = table.rows[0].cells
    header_cells[0].text = 'ชุดข้อมูล'
    header_cells[1].text = 'จำนวน Records'
    header_cells[2].text = 'คำอธิบาย'
    
    # Data
    data = [
        ('cultivation.csv', '6,226', 'ข้อมูลการเพาะปลูก'),
        ('crop_characteristics.csv', '46', 'ลักษณะพืช 46 ชนิด'),
        ('weather.csv', '56,287', 'ข้อมูลสภาพอากาศ'),
        ('price.csv', '2,289,492', 'ข้อมูลราคาผลผลิต'),
        ('economic.csv', '731', 'ข้อมูลเศรษฐกิจ'),
    ]
    
    for i, (name, count, desc) in enumerate(data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = name
        row_cells[1].text = count
        row_cells[2].text = desc
    
    doc.add_paragraph()
    
    # ภาคผนวก ค
    doc.add_heading('ภาคผนวก ค: ผลการทดสอบเพิ่มเติม', level=2)
    
    test_results = """
ผลการทดสอบ Model C (Price Forecasting):

Baseline Model:
- MAE: 3.01 บาท
- RMSE: 4.13 บาท
- Price Bias: 96.79%
- Features: 13 (temporal only)

Improved Model (Minimal Dataset):
- MAE: 13.31 บาท
- RMSE: 18.91 บาท
- Price Bias: 68.09% (ลดลง 28.7%)
- Features: 21 (13 + 8 external)
- Weather Features: 3.33%
- Economic Features: 3.33%

Improved Model (Full Dataset):
- MAE: 0.49 บาท (ดีที่สุด)
- RMSE: 0.79 บาท
- MAPE: 1.18%
- Price Bias: 99.71% (แย่กว่า - overfitting)

สรุป: Minimal Dataset Model เหมาะสมกว่าสำหรับ production เพราะมี robustness สูงกว่า
    """
    
    doc.add_paragraph(test_results.strip())

if __name__ == "__main__":
    main()
